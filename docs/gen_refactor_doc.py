# -*- coding: utf-8 -*-
"""AMORI 리팩토링 방향 제안서 생성 스크립트 (Word .docx, 네이티브 표)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KOREAN_FONT = "맑은 고딕"
ACCENT = RGBColor(0xC0, 0x4A, 0x3A)        # terracotta
HEADER_BG = "2F2A26"                        # dark
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_BG = "F4F1EA"                          # cream

doc = Document()

# ---- 기본 폰트 (라틴 + 동아시아 모두 한글 폰트로) ----
normal = doc.styles["Normal"]
normal.font.name = KOREAN_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

for sname in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
    st = doc.styles[sname]
    st.font.name = KOREAN_FONT
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_cell_text(cell, bold=False, color=None, size=9.5, align=None):
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        for r in p.runs:
            r.font.name = KOREAN_FONT
            r.font.size = Pt(size)
            r.font.bold = bold
            r._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
            if color is not None:
                r.font.color.rgb = color


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], HEADER_BG)
        style_cell_text(hdr[i], bold=True, color=HEADER_FG, size=9.5,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            if ri % 2 == 1:
                set_cell_bg(cells[ci], ZEBRA_BG)
            style_cell_text(cells[ci], size=9.5)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT


def body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(10.5)
    run.font.name = KOREAN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = KOREAN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    return p


# ============================ 표지 ============================
title = doc.add_heading("AMORI 리팩토링 방향 제안", level=0)
for r in title.runs:
    r.font.color.rgb = ACCENT
sub = doc.add_paragraph("AI 에이전트 소개팅 데이팅앱 · KT 자체모델 → 외부 LLM API 피벗 대응")
sub.runs[0].font.size = Pt(11)
sub.runs[0].font.italic = True
meta = doc.add_paragraph("작성일: 2026-06-10 · 코드 미수정, 분석 및 방향 설정 문서")
meta.runs[0].font.size = Pt(9)
meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph()

# ============================ 1 ============================
h1("1. 현재 상태 진단 — 세 개의 분리된 세계")
body("코드 전체를 읽고 내린 결론은, 지금 레포에 같은 제품의 서로 다른 구현이 3개 존재한다는 것입니다.")

doc.add_heading("① Flutter 단독 세계 (현재 실제로 돌아가는 것)", level=2)
bullet("lib/core/services/llm_service.dart:16 — Flutter가 Groq API(llama-3.1-8b-instant)를 직접 호출. "
       "페르소나 생성, 에이전트 대화 생성, 궁합 리포트까지 전부 클라이언트에서 LLM을 부름.")
bullet("결과는 PersonaStore 정적 싱글톤에 메모리로만 들고 있다가 일부를 Firestore에 씀(amori_backend.dart:90-101).")
bullet("대화 상대는 conversation_service.dart:12-19에 하드코딩된 \"민준\" 한 명뿐. "
       "대화는 한 번의 LLM 호출로 8~10개 메시지를 한꺼번에 생성. 타이핑 인디케이터는 순수 연출.")

doc.add_heading("② FastAPI 백엔드 세계 (설계는 좋은데 실행이 안 되는 것)", level=2)
bullet("라우터/스키마/SSE 스트리밍/Firebase 인증/레이트리밋까지 구조는 잘 잡혀 있음. "
       "LLM 추상화(backend/app/llm/base.py)도 4개 도메인 메서드로 깔끔.")
bullet("치명적 결함: app/models/database.py가 존재하지 않음. 모든 라우터가 이 파일을 import하는데 "
       "없어서 백엔드는 현재 기동 자체가 불가능. alembic도 같은 이유로 마이그레이션 불가.")
bullet("Flutter에는 이 백엔드를 호출하는 HTTP 클라이언트가 한 줄도 없음. 두 세계는 연결된 적이 없음.")

doc.add_heading("③ KT 시대의 유산 (전제가 사라진 것)", level=2)
bullet("llm/ 디렉토리: KT Mi:dm 2.0을 대회 GPU에 셀프호스팅하고 4개 HTTP 엔드포인트를 노출하는 별도 추론 서버 계획(README만 존재).")
bullet("backend/app/llm/midm_local.py, hf.py: 그 추론 서버를 호출하는 HTTP 어댑터.")
bullet("이 구조 전체가 \"자체 모델을 GPU에 띄워야 한다\"는 전제 위에 서 있음. "
       "외부 LLM API로 피벗한 지금, 별도 LLM 서비스 계층의 존재 이유가 사라짐.")
body("추가로 데이터 저장소도 갈라져 있음: Flutter는 Firestore에, 백엔드는 Postgres+pgvector에 쓰도록 설계되어 "
     "진실의 원천이 2개.")

# ============================ 2 ============================
h1("2. 리팩토링 핵심 결정 5가지")

doc.add_heading("결정 1 — llm/ 별도 서비스를 폐기하고 백엔드 안으로 흡수", level=2)
bullet("llm/이 별도 HTTP 서비스였던 유일한 이유는 GPU 머신에서 Mi:dm을 띄우기 위함. "
       "외부 API를 쓰면 백엔드→LLM서비스→외부API의 HTTP 홉은 지연·장애지점·배포단위만 늘리는 순수 비용.")
bullet("LLMProvider ABC가 HTTP 경로가 아닌 도메인 동작 수준에서 추상화되어 있어, 외부 API SDK를 직접 호출하는 "
       "anthropic.py provider를 추가하고 factory.py에 case 한 줄 넣으면 끝. 라우터·스키마·로직 무변경.")
bullet("프롬프트 엔지니어링은 사라지는 게 아니라 backend/app/llm/prompts/ 패키지로 이동. "
       "한국어 프롬프트 템플릿은 여전히 핵심 자산이고 담당자는 이 디렉토리만 소유하면 협업 경계도 유지.")
bullet("삭제 대상: midm_local.py(KT 전용), 셀프호스팅 계획. hf.py도 복귀 가능성 0이면 삭제. "
       "mock.py는 테스트·오프라인 데모용으로 반드시 유지(한국어 더미 대화가 이미 잘 만들어져 있음).")

doc.add_heading("결정 2 — Flutter의 직접 LLM 호출을 전부 제거하고 BFF 경유로 일원화", level=2)
body("가장 시급한 구조 변경. 근거:", bold=True)
bullet("보안: GROQ_API_KEY가 .env로 앱 번들에 들어감. Flutter 앱은 디컴파일로 키 추출이 가능하므로 출시 즉시 키 공개와 같음. "
       "(firebase_options.dart의 Firebase 키는 원래 공개 설계 값이라 문제 아님 — 진짜 방어선은 Firestore 보안 규칙.)")
bullet("비용 통제 불가: 백엔드의 일일 5회 쿼터(rate_limit.py)를 클라이언트 직접 호출이 완전히 우회. 비용이 통제 없이 증가.")
bullet("데이터 정합성: 대화 로그는 리포트·매칭·피드백 학습의 입력. 클라이언트 메모리에만 있으면 서버가 활용 불가.")
bullet("중복 제거: 페르소나/대화/리포트 로직과 한국어 프롬프트가 Flutter와 백엔드 양쪽에 존재. 개선을 두 군데서 해야 함.")
bullet("AI 기본법 대응: ai_generated:true 라벨링과 RAI 필터를 서버 한 곳에서 강제 가능.")

doc.add_heading("결정 3 — 데이터 저장소를 Postgres 단일 원천으로, Firebase는 Auth(+FCM)만", level=2)
bullet("매칭이 1024차원 벡터 코사인 유사도 기반인데 Firestore는 벡터 검색이 사실상 불가. Postgres는 어차피 필수.")
bullet("\"백엔드 스택 유지\" 요구와 일치 — FastAPI + SQLAlchemy + asyncpg + pgvector 그대로.")
bullet("클라이언트의 Firestore 직접 쓰기 경로는 제거하고 모든 도메인 데이터는 BFF 경유. "
       "두 DB에 같은 개념이 쌓이는 현 상태는 시간이 갈수록 복구 비용 증가.")
bullet("Firebase에 남기는 것: Authentication(백엔드 토큰 검증 완성됨), FCM 푸시 토큰.")

doc.add_heading("결정 4 — LLM 모델: Claude 이중 티어 추천", level=2)
body("이 앱의 제품 자체가 \"에이전트끼리의 한국어 대화\"이고 사용자가 그 대화를 직접 읽으므로, 한국어 자연스러움이 모델 선택의 1순위. "
     "현재 llama-3.1-8b-instant는 8B급으로 한국어 구어체가 어색하고 JSON 스키마 준수율도 낮아 이 용도에 가장 약한 선택. "
     "Groq의 장점은 속도뿐인데 시뮬레이션은 백그라운드라 속도가 크게 중요하지 않음.")
add_table(
    ["용도", "추천 모델", "가격 (입력/출력, 1M토큰)", "근거"],
    [
        ["시뮬레이션 턴 생성 (대량)", "claude-haiku-4-5", "$1 / $5",
         "한국어 구어체가 8B 오픈모델 대비 우수, 빠름, 최저가 티어"],
        ["페르소나 분석·궁합 리포트 (저빈도·고품질)", "claude-sonnet-4-6", "$3 / $15",
         "24문항에서 성격을 읽는 추론 품질, 리포트 문장력"],
        ["데모·발표용 (전 구간)", "claude-sonnet-4-6", "$3 / $15",
         "시연에서는 대화 품질이 곧 제품"],
    ],
    col_widths=[1.6, 1.4, 1.4, 2.6],
)
body("비용 추산: 20턴 시뮬레이션 1회 ≈ 누적 입력 ~50K + 출력 ~2K 토큰. Haiku 기준 약 $0.06, "
     "프롬프트 캐싱 적용 시 약 $0.01~0.02/회. 사용자당 일 5회 쿼터면 사용자당 일 ~10센트 이하로, "
     "MVP 100명 규모에서 부담 없음.")
body("부가 이점:", bold=True)
bullet("Structured Outputs(json_schema): shared/schemas/를 API 레벨에서 강제 → SCHEMA_VIOLATION 오류 클래스 소멸, 파싱 재시도 불필요.")
bullet("프롬프트 캐싱: 시뮬레이션 턴 루프에서 페르소나+이전 대화가 프리픽스로 반복되어 캐싱 효과 극대화.")
body("대안: Gemini 2.5 Flash·GPT-4.1-mini도 한국어 준수하고 더 저렴할 수 있음. provider 추상화를 유지하므로 "
     "나중에 gemini.py를 추가해 환경변수 스위치로 A/B 비교 가능. 처음엔 하나로 시작 권장.")
body("임베딩 (페르소나 벡터):", bold=True)
add_table(
    ["선택지", "차원", "비용/특징", "비고"],
    [
        ["Voyage AI voyage-3.5", "1024", "API, Anthropic 공식 파트너", "운영 단순, 추천"],
        ["BGE-M3 셀프호스팅", "1024", "무료, 한국어 강함, CPU 가능", "비용 0, 배포 복잡도 +1"],
        ["Cohere embed-multilingual-v3.0", "1024", "API", "대안"],
    ],
    col_widths=[2.2, 0.8, 2.4, 1.6],
)
body("Anthropic에는 임베딩 API가 없어 별도 선택 필요. 스키마가 1024차원 고정인데 위 선택지는 1024차원을 기본 출력하므로 "
     "마이그레이션 없이 그대로 사용 가능. DB 모델을 새로 쓰는 지금이 차원 변경 비용도 가장 쌈.")

doc.add_heading("결정 5 — 시뮬레이션을 \"원샷 생성\"에서 \"2-에이전트 턴 루프\"로 전환", level=2)
body("현재는 한 번의 호출로 양쪽 대사를 전부 생성. 이 방식의 품질 문제는 구조적:")
bullet("한 모델이 한 컨텍스트에서 양쪽을 쓰면 두 페르소나의 말투가 섞임(style bleed). "
       "서로 다른 성격이 5턴 만에 비슷한 어투로 수렴.")
bullet("대화가 갈등·탐색 없이 평탄한 핑퐁이 됨 — 궁합을 \"측정\"해야 하는 목적과 어긋남.")
body("제안 구조 (backend services/simulation.py):", bold=True)
bullet("에이전트 A·B 각각 자기 페르소나만 담긴 별도 시스템 프롬프트를 가진 두 개의 대화 컨텍스트 유지")
bullet("턴마다 한쪽씩 호출해 발화 생성 (Haiku, 시스템 프롬프트 캐싱)")
bullet("N턴마다 또는 종료 시 별도의 가벼운 분석 호출로 시그널·케미 점수 추출 (mock.py의 system 포맷 활용)")
bullet("각 턴을 기존 SSE 파이프라인(routers/simulation.py:31-47)으로 전송 — 이미 만든 SSE 인프라가 드디어 실데이터 운반")
body("이러면 agent_chat_screen.dart의 가짜 타이핑을 진짜 실시간 스트리밍으로 전환 가능. "
     "\"내 에이전트가 지금 소개팅 중\"이라는 제품의 핵심 감성이 연출이 아니라 실제가 됨. "
     "(Flutter SSE 소비는 http streamed request 또는 flutter_client_sse, 1차는 잡 폴링으로 시작 가능.)")

# ============================ 3 ============================
h1("3. 목표 아키텍처")
arch = (
    "Flutter 앱\n"
    " ├─ Firebase Auth (로그인, ID 토큰 발급)          ← Firebase는 여기까지만\n"
    " └─ ApiClient (Bearer ID토큰) ──→ FastAPI BFF\n"
    "                                   ├─ auth/      Firebase 토큰 검증 (현행 유지)\n"
    "                                   ├─ routers/   persona·matches·simulation(SSE)·report·meet (유지)\n"
    "                                   ├─ llm/\n"
    "                                   │   ├─ base.py        (유지 — 4개 도메인 메서드)\n"
    "                                   │   ├─ factory.py     (유지)\n"
    "                                   │   ├─ anthropic.py   ★신규: SDK 직접 호출 + structured outputs\n"
    "                                   │   ├─ mock.py        (유지 — 테스트/오프라인 데모)\n"
    "                                   │   └─ prompts/       ★신규: 한국어 프롬프트 (구 llm/ 책임 이관)\n"
    "                                   ├─ services/simulation.py  ★재작성: 2-에이전트 턴 루프\n"
    "                                   ├─ matching/  ★구 matching/ 모듈을 백엔드 내 패키지로\n"
    "                                   ├─ models/database.py  ★신규: SQLAlchemy 모델 (현재 미존재!)\n"
    "                                   └─ PostgreSQL + pgvector (단일 데이터 원천)\n"
    "                                        + 외부: Claude API / Voyage(임베딩)"
)
pre = doc.add_paragraph()
run = pre.add_run(arch)
run.font.name = "Consolas"
run.font.size = Pt(8.5)
body("matching/을 별도 서비스가 아닌 백엔드 내 패키지로 두는 근거: matching/README.md 스스로 \"라이브러리 import\" 옵션을 허용하고, "
     "4인 팀·짧은 일정에서 배포 단위는 적을수록 좋음. /matches/find 베이스라인(top-K 코사인)이 이미 백엔드에 있어 같은 프로세스 교체가 자연스러움.")
body("shared/schemas/는 계약으로 계속 유지하되 의미가 바뀜 — \"백엔드↔LLM서비스 HTTP 계약\"에서 "
     "\"LLM structured output 스키마 + 백엔드↔Flutter 응답 계약\"으로.")

# ============================ 4 ============================
h1("4. Flutter 쪽 정리 방향")
bullet("레이어 도입: data/repositories/ (PersonaRepository, SimulationRepository, MatchRepository...) → 내부에서 ApiClient 호출. 화면은 repository만 알게.")
bullet("상태 관리: 정적 싱글톤 PersonaStore(reset()이 한 번도 안 불림) 폐기, Riverpod 도입 권장(최소 변경 시 ChangeNotifier+Provider). "
       "현재는 전역 정적 필드 의존이라 세션 간 누수·null 전파(persona_loading_screen.dart:84가 LLM 실패를 삼키고 null로 진행)가 구조적으로 발생.")
bullet("UI에 박힌 비즈니스 로직 이동: persona_loading_screen.dart:68-101(LLM 3연속+저장을 initState에서), "
       "match_list_screen.dart:138-151(필터/정렬), request_status_screen.dart(카운트다운) 등 → repository로.")
bullet("하드코딩 더미와 실데이터 경로 분리: kMatches, _fallbackMessages, 하드코딩 파트너 \"민준\"은 mock provider가 백엔드에서 동일 역할 → 단계적 제거.")
bullet("즉시 청소: app_config 2.dart, debug_storage_service 2.dart (macOS 복사 충돌 중복 파일) 삭제.")
body("화면 플로우(온보딩→시나리오→로딩→홈→매칭→만남)는 잘 설계되어 UI 구조는 유지, 데이터 결선만 교체.")

# ============================ 5 ============================
h1("5. 실행 순서 제안 (의존성 기준 3단계)")
doc.add_heading("P0 — 백엔드를 살린다 (다른 모든 것의 전제)", level=2)
bullet("app/models/database.py 작성 (User, Persona, Match, SimulationJob, Report, MeetRequest, Feedback, LLMCallLog) + alembic 초기 마이그레이션 — 현재 기동 불가의 원인")
bullet("llm/anthropic.py provider + llm/prompts/ 작성 (Flutter 기존 한국어 프롬프트를 시드로 이관 — 검증된 자산)")
bullet("LLM_PROVIDER=mock→anthropic 스위치로 persona→simulation→report E2E를 백엔드 단독 검증")
bullet("midm_local.py·셀프호스팅 계획 제거, 부수 버그 수정(matches.py:67 match_id 문자열/UUID 불일치, 미사용 log_llm_call 결선)")
doc.add_heading("P1 — Flutter를 백엔드에 연결한다", level=2)
bullet("ApiClient + repository 계층, Firebase ID 토큰 헤더")
bullet("클라이언트 LLM 호출 3종 제거, persona 빌드/시뮬레이션/리포트를 BFF 경유로 교체 (1차 폴링, 2차 SSE)")
bullet("클라이언트의 Firestore 도메인 데이터 쓰기 제거")
doc.add_heading("P2 — 제품 품질을 올린다", level=2)
bullet("2-에이전트 턴 루프 + SSE 실시간 표시 (가짜 타이핑 → 진짜 스트리밍)")
bullet("임베딩 도입(Voyage 또는 BGE-M3) → /matches/find 실데이터화 → 더미 매치 4명 졸업")
bullet("matching 패키지에서 카테고리 가중치·피드백 루프 고도화, 한국어 대화 품질 평가셋(구 llm/ 계획의 30케이스) 운영")

# ============================ 6 ============================
h1("6. 유지할 것 / 버릴 것 요약")
add_table(
    ["유지 (잘 만들어짐)", "버림 (전제 소멸 / 중복)"],
    [
        ["FastAPI 라우터·스키마·SSE·인증·쿼터 구조 전체", "llm/ 별도 HTTP 서비스 계획, midm_local.py"],
        ["LLMProvider 추상화 + factory + mock.py", "Flutter의 직접 LLM 호출 3개 서비스"],
        ["shared/schemas/ 계약 (structured output 스키마로 재활용)", "클라이언트의 Firestore 도메인 쓰기 경로"],
        ["Flutter 화면 플로우·디자인 시스템(core/theme, core/widgets)", "PersonaStore 정적 싱글톤, 하드코딩 파트너 \"민준\""],
        ["Flutter의 한국어 프롬프트 (백엔드 prompts/로 이관)", "* 2.dart 중복 파일, Groq 의존"],
        ["Firebase Auth 통합 (양쪽 모두)", "1회 호출 원샷 대화 생성 방식"],
    ],
    col_widths=[3.5, 3.5],
)
body("가장 중요한 한 줄: \"LLM 호출 경로를 Flutter→Groq에서 Flutter→BFF→Claude로 일원화하고, "
     "그 전제 작업으로 존재하지 않는 DB 모델부터 작성한다\".", bold=True)

out = r"C:\git\amori\docs\AMORI_리팩토링_방향.docx"
doc.save(out)
print("SAVED:", out)
